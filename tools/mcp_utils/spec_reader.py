"""
Spec File Reader - Reads specification files in various formats.

Supports:
- Plain text (.txt, .md, .json, .yaml, .xml, .html)
- PDF documents (.pdf)
- Word documents (.docx)
- Images (.png, .jpg, .jpeg, .gif, .webp) - converted to base64 for Claude

For PDF and Word, optional dependencies are required:
- PDF: pypdf2 or pdfplumber
- Word: python-docx

If dependencies are not installed, graceful fallback with helpful error messages.
"""

import base64
import logging
import mimetypes
from pathlib import Path
from typing import Dict, Optional, Tuple

logger = logging.getLogger(__name__)

# File type categories
TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".yaml",
    ".yml",
    ".xml",
    ".html",
    ".htm",
    ".csv",
    ".rst",
    ".adoc",
    ".ini",
    ".cfg",
    ".conf",
    ".toml",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}

PDF_EXTENSIONS = {".pdf"}

WORD_EXTENSIONS = {".docx", ".doc"}


def read_spec_file(file_path: str) -> Tuple[str, str, Optional[str]]:
    """
    Read a specification file and return its content.

    Args:
        file_path: Path to the spec file

    Returns:
        Tuple of (content, content_type, error_message)
        - content: The file content (text or base64 for images)
        - content_type: "text", "image", or "error"
        - error_message: None if successful, error string if failed
    """
    path = Path(file_path)

    if not path.exists():
        return "", "error", f"File not found: {file_path}"

    ext = path.suffix.lower()

    try:
        # Plain text files
        if ext in TEXT_EXTENSIONS:
            return _read_text_file(path)

        # PDF files
        elif ext in PDF_EXTENSIONS:
            return _read_pdf_file(path)

        # Word documents
        elif ext in WORD_EXTENSIONS:
            return _read_word_file(path)

        # Images
        elif ext in IMAGE_EXTENSIONS:
            return _read_image_file(path)

        # Unknown - try as text
        else:
            logger.warning(f"Unknown file type {ext}, attempting to read as text")
            return _read_text_file(path)

    except Exception as e:
        return "", "error", f"Failed to read {file_path}: {str(e)}"


def _read_text_file(path: Path) -> Tuple[str, str, Optional[str]]:
    """Read a plain text file."""
    try:
        content = path.read_text(encoding="utf-8")
        return content, "text", None
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            content = path.read_text(encoding="latin-1")
            return content, "text", None
        except Exception as e:
            return "", "error", f"Encoding error: {str(e)}"


def _read_pdf_file(path: Path) -> Tuple[str, str, Optional[str]]:
    """Read a PDF file and extract text."""
    # Try pypdf first (newer, maintained fork of PyPDF2)
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

        if text_parts:
            return "\n\n".join(text_parts), "text", None
        else:
            return (
                "",
                "error",
                "PDF contains no extractable text (may be scanned/image-based)",
            )

    except ImportError:
        pass

    # Try PyPDF2 (older but common)
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(path))
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

        if text_parts:
            return "\n\n".join(text_parts), "text", None
        else:
            return (
                "",
                "error",
                "PDF contains no extractable text (may be scanned/image-based)",
            )

    except ImportError:
        pass

    # Try pdfplumber (best for complex PDFs)
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

        if text_parts:
            return "\n\n".join(text_parts), "text", None
        else:
            return (
                "",
                "error",
                "PDF contains no extractable text (may be scanned/image-based)",
            )

    except ImportError:
        pass

    # No PDF library available
    return (
        "",
        "error",
        (
            "PDF support requires one of: pypdf, PyPDF2, or pdfplumber\n"
            "Install with: pip install pypdf  (recommended)\n"
            "         or: pip install pdfplumber  (for complex PDFs)"
        ),
    )


def _read_word_file(path: Path) -> Tuple[str, str, Optional[str]]:
    """Read a Word document and extract text."""
    ext = path.suffix.lower()

    # .doc files (old format) - not easily supported
    if ext == ".doc":
        return (
            "",
            "error",
            (
                "Old .doc format not supported. Please convert to .docx\n"
                "(Open in Word and Save As .docx)"
            ),
        )

    # .docx files
    try:
        from docx import Document

        doc = Document(str(path))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        if text_parts:
            return "\n\n".join(text_parts), "text", None
        else:
            return "", "error", "Word document appears to be empty"

    except ImportError:
        return (
            "",
            "error",
            (
                "Word document support requires python-docx\n"
                "Install with: pip install python-docx"
            ),
        )


def _read_image_file(path: Path) -> Tuple[str, str, Optional[str]]:
    """Read an image file and convert to base64 for Claude."""
    try:
        with open(path, "rb") as f:
            image_data = f.read()

        # Get MIME type
        mime_type, _ = mimetypes.guess_type(str(path))
        if not mime_type:
            # Fallback based on extension
            ext = path.suffix.lower()
            mime_map = {
                ".png": "image/png",
                ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg",
                ".gif": "image/gif",
                ".webp": "image/webp",
                ".bmp": "image/bmp",
            }
            mime_type = mime_map.get(ext, "image/png")

        # Encode to base64
        b64_data = base64.b64encode(image_data).decode("utf-8")

        # Return in a format that can be used with Claude's vision
        # The content includes the data URI format
        content = f"data:{mime_type};base64,{b64_data}"

        return content, "image", None

    except Exception as e:
        return "", "error", f"Failed to read image: {str(e)}"


def format_spec_for_prompt(
    file_path: str,
    content: str,
    content_type: str,
) -> Dict:
    """
    Format spec content for inclusion in a prompt.

    Args:
        file_path: Original file path
        content: File content (text or base64)
        content_type: "text" or "image"

    Returns:
        Dict with formatted content ready for prompt injection
    """
    path = Path(file_path)

    if content_type == "text":
        return {
            "type": "text",
            "file_name": path.name,
            "file_path": str(path),
            "content": content,
            "formatted": (
                f"### File: {path.name}\nPath: {file_path}\n\n```\n{content}\n```"
            ),
        }
    elif content_type == "image":
        return {
            "type": "image",
            "file_name": path.name,
            "file_path": str(path),
            "content": content,  # base64 data URI
            "formatted": (
                f"### Image: {path.name}\n"
                f"Path: {file_path}\n"
                f"[Image content - Claude will analyze this visually]"
            ),
            "media_type": content.split(";")[0].replace("data:", ""),
            "base64_data": content.split(",")[1] if "," in content else content,
        }
    else:
        return {
            "type": "error",
            "file_name": path.name,
            "file_path": str(path),
            "content": "",
            "formatted": f"### Error reading: {path.name}\n{content}",
        }


def read_all_spec_files(file_paths: list) -> Dict:
    """
    Read all spec files and return combined result.

    Args:
        file_paths: List of file paths

    Returns:
        Dict with:
        - text_content: Combined text for text files
        - images: List of image data for vision
        - errors: List of error messages
        - summary: Human-readable summary
    """
    result = {
        "text_content": [],
        "images": [],
        "errors": [],
        "files_loaded": [],
        "summary": "",
    }

    for file_path in file_paths:
        content, content_type, error = read_spec_file(file_path)

        if error:
            result["errors"].append(f"{file_path}: {error}")
            continue

        formatted = format_spec_for_prompt(file_path, content, content_type)

        if content_type == "text":
            result["text_content"].append(formatted["formatted"])
            result["files_loaded"].append(f"✓ {Path(file_path).name} (text)")

        elif content_type == "image":
            result["images"].append(
                {
                    "file_name": formatted["file_name"],
                    "media_type": formatted["media_type"],
                    "base64_data": formatted["base64_data"],
                }
            )
            result["text_content"].append(formatted["formatted"])
            result["files_loaded"].append(f"✓ {Path(file_path).name} (image)")

    # Build summary
    summary_parts = []
    if result["files_loaded"]:
        summary_parts.append(f"Loaded {len(result['files_loaded'])} spec file(s):")
        summary_parts.extend(f"  {f}" for f in result["files_loaded"])

    if result["errors"]:
        summary_parts.append(f"\nErrors ({len(result['errors'])}):")
        summary_parts.extend(f"  ⚠ {e}" for e in result["errors"])

    result["summary"] = "\n".join(summary_parts)

    return result
